from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

# ── Title ──────────────────────────────────────────────────────────────────────
title = doc.add_heading("Composing the Browser URL for a GAS Application", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 1. Overview ───────────────────────────────────────────────────────────────
doc.add_heading("1. Overview", level=2)
doc.add_paragraph(
    "When a Genero application is deployed on the Genero Application Server (GAS) "
    "and exposed through Apache HTTP Server, the browser URL that launches the "
    "application is built by combining several independent pieces of information "
    "drawn from different parts of the infrastructure. Understanding each segment "
    "is essential for correctly constructing or troubleshooting application URLs."
)

# ── 2. URL anatomy ─────────────────────────────────────────────────────────────
doc.add_heading("2. URL Structure", level=2)
doc.add_paragraph("The full URL has the following form:")

url_para = doc.add_paragraph(style="No Spacing")
run = url_para.add_run("    http://<host>:<port>/genero/<type>/r/<group>/<appname>")
run.font.name = "Courier New"
run.font.size = Pt(11)
run.font.bold = True
doc.add_paragraph()  # spacer

table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Segment"
hdr[1].text = "Description"
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True

rows = [
    ("http or https",
     "Protocol. Determined by the Apache virtual-host configuration (HTTP on port 80, "
     "HTTPS on port 443). Not stored in the GAS configuration."),
    ("<host>",
     "Hostname or IP address of the server running Apache. Typically 'localhost' for "
     "local access or the server's public DNS name for remote access."),
    (":<port>",
     "Port on which Apache listens. Omitted when using the default ports (80/443). "
     "In development setups Apache may listen on a non-standard port."),
    ("/genero",
     "The Apache alias that maps the URL path to the GAS FastCGI dispatcher. "
     "This segment is defined entirely in the Apache configuration and is NOT part "
     "of the GAS configuration. See Section 4 for full details."),
    ("<type>",
     "Routing type. 'ua' for user-agent (interactive) applications; "
     "'ws' for web-service applications. Derived from the APPLICATION_GROUPS or "
     "SERVICE_GROUPS XML file."),
    ("/r/",
     "Fixed literal segment used by the GAS dispatcher to identify a routing request."),
    ("<group>",
     "The group identifier (Id attribute) from the APPLICATION_GROUPS or SERVICE_GROUPS "
     "XML file whose path matches the directory containing the application XCF file."),
    ("<appname>",
     "The application name, which is the XCF filename without its .xcf extension."),
]

for seg, desc in rows:
    row = table.add_row().cells
    row[0].text = seg
    run = row[0].paragraphs[0].runs[0]
    run.font.name = "Courier New"
    row[1].text = desc

doc.add_paragraph()

# ── 3. Where each segment comes from ──────────────────────────────────────────
doc.add_heading("3. Where Each Segment Is Defined", level=2)

doc.add_heading("3.1  Port", level=3)
doc.add_paragraph(
    "The GAS server port is defined in the as.xcf configuration file via the resource "
    "res.ic.server.port, which defaults to 6394. This value is referenced in the "
    "INTERFACE_TO_CONNECTOR section:"
)
p = doc.add_paragraph(style="No Spacing")
r = p.add_run(
    '    <RESOURCE Id="res.ic.server.port" Source="INTERNAL">6394</RESOURCE>\n'
    '    ...\n'
    '    <TCP_SERVER_PORT>$(res.ic.server.port)</TCP_SERVER_PORT>'
)
r.font.name = "Courier New"
r.font.size = Pt(10)
doc.add_paragraph()

doc.add_paragraph(
    "Note: this is the port on which the GAS dispatcher's built-in TCP listener "
    "accepts FastCGI connections from Apache — it is NOT the port the browser connects "
    "to when Apache is in front."
)

doc.add_heading("3.2  Application Group and Type (ua / ws)", level=3)
doc.add_paragraph(
    "The group identifier and routing type (ua or ws) come from XML group files. "
    "Two types exist:"
)
doc.add_paragraph(
    "APPLICATION_GROUPS — defines groups for interactive (ua) applications. "
    "Each GROUP element maps an Id to a filesystem directory path. "
    "If the application's XCF file resides in that directory, the group Id is used.",
    style="List Bullet"
)
doc.add_paragraph(
    "SERVICE_GROUPS — same structure, but for web-service (ws) applications.",
    style="List Bullet"
)
doc.add_paragraph("Example APPLICATION_GROUPS file:")
p = doc.add_paragraph(style="No Spacing")
r = p.add_run(
    '    <APPLICATION_GROUPS>\n'
    '        <GROUP Id="dev">/home/genero/groups/apps</GROUP>\n'
    '    </APPLICATION_GROUPS>'
)
r.font.name = "Courier New"
r.font.size = Pt(10)
doc.add_paragraph()
doc.add_paragraph(
    "An application whose XCF file is stored in /home/genero/groups/apps/ will receive "
    "group Id 'dev' and type 'ua', producing the path segment /ua/r/dev/ in the URL."
)
doc.add_paragraph(
    "The running fastcgidispatch process is started with --application-group and "
    "--service-group flags pointing to these files. The xcfinfo tool discovers "
    "the active group files by inspecting the running process's command line."
)

doc.add_heading("3.3  Application Name", level=3)
doc.add_paragraph(
    "The application name is simply the base filename of the XCF file with the "
    ".xcf extension removed. For example, ifx1.xcf produces the name ifx1."
)

# ── 4. Apache and the /genero prefix ──────────────────────────────────────────
doc.add_heading("4. Why Apache Adds the “genero” Prefix", level=2)

doc.add_heading("4.1  The Role of Apache", level=3)
doc.add_paragraph(
    "The /genero path segment does not exist in any GAS configuration file (as.xcf, "
    "group XML files, or application XCF files). It is not a GAS concept at all. "
    "It is the URL path alias defined in the Apache HTTP Server configuration that "
    "maps incoming browser requests to the GAS FastCGI dispatcher binary."
)
doc.add_paragraph(
    "Without Apache in front of the GAS dispatcher, the /genero segment would not "
    "appear in the URL. In a pure standalone GAS deployment (no Apache), the dispatcher "
    "is accessed directly and there is no such prefix."
)

doc.add_heading("4.2  The Apache Configuration", level=3)
doc.add_paragraph(
    "On this installation the Genero Apache configuration is located at:"
)
p = doc.add_paragraph(style="No Spacing")
r = p.add_run("    /etc/apache2/conf-available/genero.conf")
r.font.name = "Courier New"
r.font.size = Pt(10)
doc.add_paragraph()
doc.add_paragraph("Its content is:")
p = doc.add_paragraph(style="No Spacing")
r = p.add_run(
    '    <IfModule mod_proxy_fcgi.c>\n'
    '      SetEnvIf Request_URI . proxy-fcgi-pathinfo=unescape\n'
    '      RewriteEngine on\n'
    '      RewriteRule .* - [E=HTTP_AUTHORIZATION:%{HTTP:Authorization}]\n\n'
    '      ProxyPass /genero/ fcgi://localhost:6394/ enablereuse=on timeout=100\n'
    '      Alias /genero /opt/genero/gas/bin/fastcgidispatch\n'
    '    </IfModule>'
)
r.font.name = "Courier New"
r.font.size = Pt(10)
doc.add_paragraph()

doc.add_heading("4.3  What Each Directive Does", level=3)
doc.add_paragraph(
    "There are two directives that together make /genero work:"
)

doc.add_heading("ProxyPass /genero/ fcgi://localhost:6394/", level=4)
doc.add_paragraph(
    "This is the operative directive. It tells Apache's mod_proxy_fcgi module to "
    "forward any HTTP request whose URL path begins with /genero/ to the FastCGI "
    "listener running on localhost port 6394 — which is the GAS dispatcher. "
    "The path prefix /genero/ is stripped when forwarding: a browser request for "
    "/genero/ua/r/dev/ifx1 becomes /ua/r/dev/ifx1 as seen by the GAS dispatcher. "
    "The enablereuse=on option keeps FastCGI connections alive for performance, "
    "and timeout=100 sets the upstream response timeout in seconds."
)

doc.add_heading("Alias /genero /opt/genero/gas/bin/fastcgidispatch", level=4)
doc.add_paragraph(
    "This directive is present for compatibility with the older mod_fcgid or "
    "CGI-based deployment model, where Apache executes the dispatcher binary directly "
    "rather than proxying to a running process. In a modern mod_proxy_fcgi setup "
    "(as above), the ProxyPass directive takes precedence and this Alias is not used "
    "for active traffic. It is kept for fallback purposes."
)

doc.add_heading("SetEnvIf and RewriteRule", level=4)
doc.add_paragraph(
    "SetEnvIf Request_URI . proxy-fcgi-pathinfo=unescape instructs mod_proxy_fcgi "
    "to decode PATH_INFO correctly, which is required for GAS routing to work. "
    "The RewriteRule preserves the HTTP Authorization header across the proxy "
    "boundary, which is needed for applications using HTTP-level authentication."
)

doc.add_heading("4.4  Summary: How /genero Ends Up in the URL", level=3)
doc.add_paragraph(
    "The name 'genero' is an arbitrary but conventional alias name chosen when "
    "configuring Apache for a Genero installation. It could in principle be any "
    "string (e.g. /gas, /app, /myserver), but /genero is the name used in all "
    "standard Genero Apache configuration examples and is therefore the de-facto "
    "standard. The GAS software itself has no knowledge of this name."
)
doc.add_paragraph(
    "When a browser navigates to http://host/genero/ua/r/dev/ifx1, the following "
    "chain occurs:"
)
steps = [
    "Apache receives the request on port 80 (or 443 for HTTPS).",
    "mod_proxy_fcgi matches the /genero/ prefix via the ProxyPass directive.",
    "Apache forwards the request as a FastCGI call to localhost:6394, stripping "
    "the /genero prefix.",
    "The GAS dispatcher receives /ua/r/dev/ifx1 and routes it to the 'dev' group.",
    "The GAS locates the ifx1.42m module in the group's directory and launches "
    "the DVM to execute the application.",
    "The response travels back through Apache to the browser.",
]
for step in steps:
    doc.add_paragraph(step, style="List Number")

# ── 5. Putting it all together ─────────────────────────────────────────────────
doc.add_heading("5. Putting It All Together — Worked Example", level=2)
doc.add_paragraph(
    "Given the following environment:"
)
env_items = [
    "Application XCF file: /home/genero/groups/apps/ifx1.xcf",
    "APPLICATION_GROUPS file: GROUP Id='dev' → /home/genero/groups/apps",
    "GAS port (res.ic.server.port): 6394",
    "Apache alias: /genero → fastcgidispatch (fcgi://localhost:6394/)",
    "Apache listening on: port 80 (HTTP)",
]
for item in env_items:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph("The resulting browser URL is:")
p = doc.add_paragraph(style="No Spacing")
r = p.add_run("    http://localhost/genero/ua/r/dev/ifx1")
r.font.name = "Courier New"
r.font.size = Pt(11)
r.font.bold = True
doc.add_paragraph()
doc.add_paragraph(
    "If Apache is configured on a non-standard port (e.g. 8080), the port appears "
    "explicitly: http://localhost:8080/genero/ua/r/dev/ifx1."
)
doc.add_paragraph(
    "Note: the GAS port 6394 never appears in the browser URL when Apache is in front. "
    "It is an internal FastCGI communication port between Apache and the GAS dispatcher."
)

# ── AI Generated footer ────────────────────────────────────────────────────────
doc.add_paragraph()
footer_para = doc.add_paragraph()
footer_run = footer_para.add_run(
    f"AI Generated - claude-sonnet-4-6 - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
)
footer_run.font.size = Pt(8)
footer_run.font.italic = True
footer_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.save("/home/genero/work/tests/xcfInfo/gas_url_composition.docx")
print("Done: gas_url_composition.docx")
